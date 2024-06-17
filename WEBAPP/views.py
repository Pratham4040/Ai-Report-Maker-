from flask import Blueprint ,render_template, request , flash,redirect,url_for,send_file
import google.generativeai as genai # type: ignore
from  wtforms import FileField
from flask_wtf import FlaskForm
import os, shutil


from WEBAPP import create_app
apikey = "API KEY HERE"
genai.configure(api_key=apikey)

generation_config = {
  "temperature": 1,
  "top_p": 0.95,
  "top_k": 64,
  "max_output_tokens": 8192,
  "response_mime_type": "text/plain",
}
safety_settings = [
  {
    "category": "HARM_CATEGORY_HARASSMENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE",
  },
  {
    "category": "HARM_CATEGORY_HATE_SPEECH",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE",
  },
  {
    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE",
  },
  {
    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE",
  },
]
model = genai.GenerativeModel(
  model_name="gemini-1.5-flash-latest",
  safety_settings=safety_settings,
  generation_config=generation_config,
)

views = Blueprint('views',__name__)
@views.route('/', methods=['POST','GET'])
def home():
        if request.method == 'POST':
            if os.path.isfile("C:\ai report maker\WEBAPP\a.docx"):
              os.remove("C:\ai report maker\WEBAPP\a.docx")
            global aim 
            aim = request.form.get('aim')
            global description 
            description = request.form.get('description')
            global material_required
            material_required = request.form.get('material_required')
            global Experiment_Title
            Experiment_Title = request.form.get('Experiment_Title')
            print(aim)
            print(description)
            print(material_required)
            print(Experiment_Title)

            first = str("aim [")
            after_aim = str("] Description [")
            after_description = str("] the materials that were required were [")
            # Prompt for aim
            after_material_foraim = str("] make a project report containing sections as Aim:(just copy down whatever aim is given by me),Materials Required(List down the materials given by me),code snippid(analyse the description and aim to give the suitable code),Working (pointwise explain the Working of the project in a (brief way no need to explain it in deep),procedure(step by step process to setup and perform the experiment(pointwise 150 words in total)),result(A suitable result for the project (50 words)),learning(what the students learned (50 words)),Right now only give me the aim without the heading")
            # Prompt for material req
            materialsrequired = str("] make a project report containing sections as Aim:(just copy down whatever aim is given by me),Materials Required(List down the materials given by me),code snippid(analyse the description and aim to give the suitable code),Working (pointwise explain the Working of the project in a (brief way no need to explain it in deep),procedure(step by step process to setup and perform the experiment(pointwise 150 words in total)),result(A suitable result for the project (50 words)),learning(what the students learned (50 words)),Right now only give me the Materials Required without the heading")
            # Prompt for code snippid
            code_snippid =str("] make a project report containing sections as Aim:(just copy down whatever aim is given by me),Materials Required(List down the materials given by me),code snippid(analyse the description and aim to give the suitable code),Working (pointwise explain the Working of the project in a (brief way no need to explain it in deep),procedure(step by step process to setup and perform the experiment(pointwise 150 words in total)),result(A suitable result for the project (50 words)),learning(what the students learned (50 words)),Right now only give me the Code Snippid without the heading")
            # Prompt for working
            working = str("] make a project report containing sections as Aim:(just copy down whatever aim is given by me),Materials Required(List down the materials given by me),code snippid(analyse the description and aim to give the suitable code),Working (pointwise explain the Working of the project in a (brief way no need to explain it in deep),procedure(step by step process to setup and perform the experiment(pointwise 150 words in total)),result(A suitable result for the project (50 words)),learning(what the students learned (50 words)),Right now only give me the Working without the heading")
            # Prompt for procedure
            procedure =str("] make a project report containing sections as Aim:(just copy down whatever aim is given by me),Materials Required(List down the materials given by me),code snippid(analyse the description and aim to give the suitable code),Working (pointwise explain the Working of the project in a (brief way no need to explain it in deep),procedure(step by step process to setup and perform the experiment(pointwise 150 words in total)),result(A suitable result for the project (50 words)),learning(what the students learned (50 words)),Right now only give me the Procedure without the heading")
            # Prompt for result
            result = str("] make a project report containing sections as Aim:(just copy down whatever aim is given by me),Materials Required(List down the materials given by me),code snippid(analyse the description and aim to give the suitable code),Working (pointwise explain the Working of the project in a (brief way no need to explain it in deep),procedure(step by step process to setup and perform the experiment(pointwise 150 words in total)),result(A suitable result for the project (50 words)),learning(what the students learned (50 words)),Right now only give me the Results without the heading")
            # Prompt for learning
            learning = str("] make a project report containing sections as Aim:(just copy down whatever aim is given by me),Materials Required(List down the materials given by me),code snippid(analyse the description and aim to give the suitable code),Working (pointwise explain the Working of the project in a (brief way no need to explain it in deep),procedure(step by step process to setup and perform the experiment(pointwise 150 words in total)),result(A suitable result for the project (50 words)),learning(what the students learned (50 words)),Right now only give me the learning without the heading" )
            #putting together all the prompts
            The_aimprompt = (first + aim + after_aim + description +after_description+material_required+after_material_foraim)
            The_materialprompt = first + aim + after_aim + description +after_description+material_required+materialsrequired
            The_codeprompt = first + aim + after_aim + description +after_description+material_required+code_snippid
            The_workingprompt = first + aim + after_aim + description +after_description+material_required+working
            The_procedurprompt = first + aim + after_aim + description +after_description+material_required+procedure
            The_resultprompt = first + aim + after_aim + description +after_description+material_required+result
            The_learningprompt = first + aim + after_aim + description +after_description+material_required+learning
            #starting the AI
            chat_session = model.start_chat()
            response = chat_session.send_message(The_aimprompt)
            global AimbyAI 
            AimbyAI = response.text
            #print(AimbyAI)
            response = chat_session.send_message(The_materialprompt)
            global MaterialsbyAI 
            MaterialsbyAI = response.text
            #print(MaterialsbyAI)
            response = chat_session.send_message(The_codeprompt)
            global CodebyAI
            CodebyAI = response.text
            #print(CodebyAI)
            response = chat_session.send_message(The_workingprompt)
            global WorkingbyAI
            WorkingbyAI = response.text
            #print(WorkingbyAI)
            response = chat_session.send_message(The_procedurprompt)
            global ProcedurebyAI
            ProcedurebyAI = response.text
            #print(ProcedurebyAI)
            response = chat_session.send_message(The_resultprompt)
            global ResultbyAI
            ResultbyAI = response.text
            #print(ResultbyAI)
            response = chat_session.send_message(The_learningprompt)
            global LearningbyAI
            LearningbyAI = response.text
            print(LearningbyAI)
            # MAKING A DOCUMENT
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH 
            from docx.oxml.ns import qn 
            Document = Document()
            Style = Document.styles['Normal']
            Style.font.name = 'Times New Roman'
            Style.font.size = Pt(14)
            Heading = Document.add_heading(Experiment_Title, 0)
            tittle_style = Heading.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(26)
            Heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rFonts = tittle_style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), "Times New Roman")
            # FOR AIM
            HAim=Document.add_heading("Aim",1)
            tittle_style = HAim.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HAim.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Para = Document.add_paragraph(AimbyAI)
            Para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            #for MATERIALS REQUIRED
            HMatreq=Document.add_heading("Material Required",1)
            tittle_style = HMatreq.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HMatreq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Para = Document.add_paragraph(MaterialsbyAI)
            Para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            #for CODE SNIPPID
            HCode=Document.add_heading("Code Snippet",1)
            tittle_style = HCode.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HCode.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Para = Document.add_paragraph(CodebyAI)
            Para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            #for WORKING
            HWork=Document.add_heading("Working",1)
            tittle_style = HWork.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HWork.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Para = Document.add_paragraph(WorkingbyAI)
            Para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #for PROCEDURE
            HProc=Document.add_heading("Procedure",1)
            tittle_style = HProc.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HProc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Para = Document.add_paragraph(ProcedurebyAI)
            Para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #for RESULT
            HRes=Document.add_heading("Result",1)
            tittle_style = HRes.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HRes.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Para = Document.add_paragraph(ResultbyAI)
            Para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #for LEARNING
            HLearn=Document.add_heading("Learning",1)
            tittle_style = HLearn.style
            tittle_style.font.name = 'Times New Roman'
            tittle_style.font.size = Pt(16)
            HLearn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Para = Document.add_paragraph(LearningbyAI)
            Para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            filepath = "c:\\ai report maker\\WEBAPP\\a.docx"
            Document.save(filepath)            
            return render_template("result.html")
        return render_template("home.html")


@views.route('/result', methods=['POST','GET'])
def result():
  return render_template("result.html")


@views.route('/download')
def download():
        path =  "a.docx"
        return send_file(path, as_attachment=True)       